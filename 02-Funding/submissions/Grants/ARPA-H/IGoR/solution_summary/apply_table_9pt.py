"""
Post-process the docx to:
1. Set all table cell text to 9pt Inter Variable (allowed per spec).
2. Tighten table cell paragraph spacing to 1pt before/after.
3. Add a light header tint to first row of each table (#EDE7F6).
4. Reduce figure image to ~4.5in width to save vertical space.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, re

SRC = "/home/mohammadi/Claude/Projects/Grants/02-submissions/Grants/ARPA-H/IGoR/solution_summary/IGoR_Solution_Summary_SUBMISSION_2026-06-19.docx"
DST = SRC  # overwrite in place

TINT_COLOR = "EDE7F6"  # light lavender header fill

def set_cell_shading(cell, fill_hex):
    """Set background shading on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    # Replace existing shd if present
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

doc = Document(SRC)

for tbl_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)
        for cell in row.cells:
            # Set shading on header row
            if is_header:
                set_cell_shading(cell, TINT_COLOR)
            # Set font and spacing on all paragraphs in cell
            for para in cell.paragraphs:
                # Paragraph spacing
                pf = para.paragraph_format
                pf.space_before = Pt(1)
                pf.space_after  = Pt(1)
                pf.line_spacing = Pt(11)
                # Font on runs
                for run in para.runs:
                    run.font.name = "Inter Variable"
                    run.font.size = Pt(9)
                # Also set via paragraph style font (direct formatting wins)
                if para.style and para.style.font:
                    try:
                        para.style.font.size = Pt(9)
                    except Exception:
                        pass
                # Apply via rPr default for the paragraph (rPrChange element not needed)
                # Set pPr/rPr directly on the paragraph xml for cells with no runs
                pPr = para._p.get_or_add_pPr()
                rPr = pPr.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    pPr.append(rPr)
                sz = rPr.find(qn("w:sz"))
                if sz is None:
                    sz = OxmlElement("w:sz")
                    rPr.append(sz)
                sz.set(qn("w:val"), "18")   # half-points: 18 = 9pt
                szCs = rPr.find(qn("w:szCs"))
                if szCs is None:
                    szCs = OxmlElement("w:szCs")
                    rPr.append(szCs)
                szCs.set(qn("w:val"), "18")

# Shrink images in the document to at most 4.5in wide
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        pass  # Can't resize via rels; need to find inline shapes

# Resize all inline images
for para in doc.paragraphs:
    for run in para.runs:
        # Check for drawing elements
        for drawing in run._r.findall('.//' + qn('wp:extent')):
            cx = int(drawing.get('cx', 0))
            cy = int(drawing.get('cy', 0))
            max_cx = int(4.5 * 914400)  # 4.5 inches in EMUs
            if cx > max_cx and cx > 0:
                ratio = max_cx / cx
                drawing.set('cx', str(max_cx))
                drawing.set('cy', str(int(cy * ratio)))
        # Also handle inline drawing extents differently
        for docPr in run._r.findall('.//' + qn('a:ext')):
            cx = int(docPr.get('cx', 0))
            cy = int(docPr.get('cy', 0))
            max_cx = int(4.5 * 914400)
            if cx > max_cx and cx > 0:
                ratio = max_cx / cx
                docPr.set('cx', str(max_cx))
                docPr.set('cy', str(int(cy * ratio)))

doc.save(DST)
print(f"Post-processed and saved: {DST}")
