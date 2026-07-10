"""
Post-process docx:
1. Set table column widths explicitly (for the team table in S4).
2. Apply 9pt font, tighten spacing, and add header tint to all tables.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/mohammadi/Claude/Projects/Grants/02-submissions/Grants/ARPA-H/IGoR/solution_summary/IGoR_Solution_Summary_SUBMISSION_2026-06-19.docx"
DST = SRC

TINT_COLOR = "EDE7F6"

def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_col_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    twips = int(width_inches * 1440)
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")

doc = Document(SRC)

# Page text width: 8.5 - 2.0 = 6.5 inches = 9360 twips
TEXT_WIDTH = 6.5  # inches

for tbl_idx, table in enumerate(doc.tables):
    ncols = len(table.columns)

    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)

        # Set column widths for 3-column tables (team table)
        if ncols == 3:
            # Focus area / role | Member and capability | Status
            # 1.5 | 3.8 | 1.2 = 6.5
            col_widths = [1.5, 3.8, 1.2]
        elif ncols == 4:
            # Generic 4-col: equal distribution
            col_widths = [TEXT_WIDTH / 4] * 4
        elif ncols == 5:
            col_widths = [1.1, 1.5, 1.2, 1.2, 1.5]
        elif ncols == 6:
            col_widths = [1.4, 1.0, 1.0, 1.0, 1.0, 1.1]
        else:
            col_widths = [TEXT_WIDTH / ncols] * ncols

        for col_idx, cell in enumerate(row.cells):
            if is_header:
                set_cell_shading(cell, TINT_COLOR)

            if col_idx < len(col_widths):
                set_col_width(cell, col_widths[col_idx])

            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.space_before = Pt(1)
                pf.space_after  = Pt(1)
                pf.line_spacing = Pt(11)

                for run in para.runs:
                    run.font.name = "Inter Variable"
                    run.font.size = Pt(9)

                pPr = para._p.get_or_add_pPr()
                rPr = pPr.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    pPr.append(rPr)
                for tag in ["w:sz", "w:szCs"]:
                    el = rPr.find(qn(tag))
                    if el is None:
                        el = OxmlElement(tag)
                        rPr.append(el)
                    el.set(qn("w:val"), "18")

doc.save(DST)
print(f"Saved: {DST}")
